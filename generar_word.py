from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── Estilos ──────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Portada ──────────────────────────────────────────────
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('EXAMEN TÉCNICO')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Optimización y API del Modelo CLV')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Cálculo de TEA basado en ROA\nCustomer Lifetime Value')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Tiempo: 24 horas')
run.bold = True
run.font.size = Pt(14)

doc.add_page_break()

# ── Tabla de contenido manual ──────────────────────────
doc.add_heading('Índice', level=1)
items = [
    '1. Objetivo del examen',
    '2. Código entregado',
    '3. Tareas a realizar',
    '    3.1. Optimización y refactorización',
    '    3.2. API REST',
    '4. Requerimientos técnicos',
    '5. Criterios de evaluación',
    '6. Estructura de entrega esperada',
    '7. Anexo: resumen del modelo actual',
]
for item in items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ── 1. Objetivo ────────────────────────────────────────
doc.add_heading('1. Objetivo del examen', level=1)
doc.add_paragraph(
    'El objetivo del presente examen es evaluar la capacidad del candidato para:'
)
bullets = [
    'Analizar y comprender un código legacy escrito deliberadamente con malas prácticas.',
    'Reducir el tiempo de ejecución del modelo CLV (Customer Lifetime Value).',
    'Refactorizar la arquitectura aplicando principios de diseño y buenas prácticas.',
    'Construir una API REST funcional que exponga el modelo como servicio.',
    'Mantener la consistencia de resultados respecto a los tests proporcionados.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# ── 2. Código entregado ────────────────────────────────
doc.add_heading('2. Código entregado', level=1)
doc.add_paragraph(
    'Se entrega un proyecto Python compuesto por los siguientes archivos:'
)

table = doc.add_table(rows=14, cols=2)
table.style = 'Light Shading Accent 1'

headers = ['Archivo', 'Propósito']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for paragraph in table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

files = [
    ('parametros final.csv', 'Input con 10 operaciones (id, producto, moneda, monto, plazo, ROA target)'),
    ('costos final_1.csv', 'Costos de originación y mantenimiento por producto + moneda'),
    ('transferencia final_final.csv', 'Curvas de tasa de fondeo por plazo'),
    ('pd final_final_2.csv', 'Probabilidad de default marginal (36 meses)'),
    ('data loadeador final.py', 'DataLoader: lectura de CSVs'),
    ('costos final_final_2.py', 'CostoHandler: gestión de costos'),
    ('tasas final_1.py', 'TasaCurva: interpolación de tasas de fondeo'),
    ('pd curvas final_final.py', 'PDCurva: cálculo de supervivencia'),
    ('amortizacion final.py', 'AmortizacionFrancesa: cuota constante'),
    ('engine final_CLV.py', 'CLVEngine: cálculo de CLV y curvas'),
    ('buscador tea final.py', 'BuscadorTEA: optimización ineficiente (barrido + bisección)'),
    ('validador final final.py', 'Validador: clase inútil que solo printea warnings'),
    ('wrapper final_final_CLV.py', 'CLVWrapper: orquestador del proceso'),
]
for i, (name, desc) in enumerate(files):
    table.rows[i + 1].cells[0].text = name
    table.rows[i + 1].cells[1].text = desc

doc.add_paragraph()
doc.add_paragraph(
    'Archivo principal: clv final final.py — entry point que importa los módulos'
    ' y ejecuta el proceso completo.'
)
doc.add_paragraph(
    'Tests: test_clv final final verdadero.py — 15 tests unitarios que validan'
    ' el comportamiento del modelo.'
)

doc.add_page_break()

# ── 3. Tareas ──────────────────────────────────────────
doc.add_heading('3. Tareas a realizar', level=1)

# 3.1
doc.add_heading('3.1. Optimización y refactorización', level=2)
doc.add_paragraph(
    'Se debe reescribir el código actual eliminando las malas prácticas e'
    ' implementando un diseño limpio, eficiente y mantenible. Las tareas específicas son:'
)

tasks_opt = [
    'Eliminar variables globales y dependencias ocultas.',
    'Reemplazar todos los iterrows() por operaciones vectorizadas con pandas/numpy.',
    'Eliminar código duplicado (métodos copiados, cálculos repetidos).',
    'Implementar un algoritmo de búsqueda de raíz eficiente (Newton-Raphson,'
    ' Brent, o similar) en lugar del barrido lineal + bisección.',
    'Reducir los nested loops internos que recalculan supervivencia en cada iteración.',
    'Separar correctamente responsabilidades (Single Responsibility).',
    'Agregar type hints y docstrings.',
    'Renombrar archivos, clases, métodos y variables con nombres descriptivos.',
    'Eliminar la clase Validador o reemplazarla por validaciones reales con excepciones.',
    'Consolidar los 9 módulos en una arquitectura coherente (mínimo 3 capas: datos, lógica, presentación/API).',
    'Mantener compatibilidad con los CSVs de entrada actuales.',
]
for t in tasks_opt:
    doc.add_paragraph(t, style='List Number')

# 3.2
doc.add_heading('3.2. API REST', level=2)
doc.add_paragraph(
    'Construir una API REST utilizando FastAPI o Flask que exponga el modelo CLV. La API debe incluir:'
)

tasks_api = [
    'POST /calcular — Recibe un payload JSON con los parámetros de entrada'
    ' (producto, moneda, monto, plazo, roa_target) y devuelve la TEA óptima'
    ' junto con las curvas intermedias.',
    'POST /calcular/lote — Recibe un array de operaciones y devuelve resultados'
    ' para todas en una sola ejecución.',
    'GET /health — Endpoint de salud.',
    'Documentación automática con Swagger/OpenAPI (FastAPI lo provee nativamente).',
    'Manejo de errores con códigos HTTP apropiados (400, 422, 500).',
    'Validación de esquemas con Pydantic.',
    'Los resultados deben ser consistentes con los tests proporcionados'
    ' (tolerancia de 0.01 en TEA y CLV).',
]
for t in tasks_api:
    doc.add_paragraph(t, style='List Number')

doc.add_page_break()

# ── 4. Requerimientos técnicos ─────────────────────────
doc.add_heading('4. Requerimientos técnicos', level=1)

reqs = [
    ('Python 3.10+', 'El código debe ser compatible con Python 3.10 o superior.'),
    ('Pandas / NumPy', 'Se permite el uso de pandas y numpy para operaciones vectorizadas.'),
    ('FastAPI o Flask', 'Framework para la API REST (recomendado FastAPI por rendimiento y Swagger integrado).'),
    ('Uvicorn', 'Servidor ASGI para FastAPI.'),
    ('Pydantic', 'Validación de esquemas de entrada/salida.'),
    ('Tests', 'Los tests existentes deben seguir pasando (o adaptarse). Los resultados numéricos deben coincidir con tolerancia ≤ 0.01.'),
    ('Tiempo de ejecución', 'El proceso completo de 10 operaciones debe ejecutarse en menos de 60 segundos (actual: ~200s).'),
    ('Arquitectura', 'Separación en capas: datos → lógica de negocio → API.'),
]

table2 = doc.add_table(rows=len(reqs) + 1, cols=2)
table2.style = 'Light Shading Accent 1'
table2.rows[0].cells[0].text = 'Requerimiento'
table2.rows[0].cells[1].text = 'Detalle'
for paragraph in table2.rows[0].cells[0].paragraphs:
    for run in paragraph.runs:
        run.bold = True
for paragraph in table2.rows[0].cells[1].paragraphs:
    for run in paragraph.runs:
        run.bold = True

for i, (req, detail) in enumerate(reqs):
    table2.rows[i + 1].cells[0].text = req
    table2.rows[i + 1].cells[1].text = detail

doc.add_page_break()

# ── 5. Criterios de evaluación ─────────────────────────
doc.add_heading('5. Criterios de evaluación', level=1)

criteria = [
    ('Reducción de tiempo de ejecución', '30%'),
    ('    Tiempo total < 60s para 10 operaciones', '20%'),
    ('Calidad del código refactorizado', '20%'),
    ('    Eliminación de malas prácticas, código limpio, type hints, docstrings', ''),
    ('Funcionalidad y consistencia de la API', '15%'),
    ('    Endpoints funcionales, documentación Swagger, validación Pydantic', ''),
    ('Mantenimiento de resultados', '10%'),
    ('    Los outputs coinciden con los tests existentes (tolerancia ≤ 0.01)', ''),
    ('Arquitectura y diseño', '5%'),
    ('    Separación de capas, patrones de diseño, inyección de dependencias', ''),
]

table3 = doc.add_table(rows=len(criteria) + 1, cols=2)
table3.style = 'Light Shading Accent 1'
table3.rows[0].cells[0].text = 'Criterio'
table3.rows[0].cells[1].text = 'Peso'
for paragraph in table3.rows[0].cells[0].paragraphs:
    for run in paragraph.runs:
        run.bold = True
for paragraph in table3.rows[0].cells[1].paragraphs:
    for run in paragraph.runs:
        run.bold = True

for i, (crit, weight) in enumerate(criteria):
    table3.rows[i + 1].cells[0].text = crit
    table3.rows[i + 1].cells[1].text = weight

doc.add_page_break()

# ── 6. Estructura de entrega ───────────────────────────
doc.add_heading('6. Estructura de entrega esperada', level=1)
doc.add_paragraph(
    'El proyecto final debe tener la siguiente estructura (o similar):'
)

code_block = """proyecto_optimizado/
├── README.md
├── requirements.txt
├── data/
│   ├── parametros.csv
│   ├── costos.csv
│   ├── transferencia.csv
│   └── pd.csv
├── clv/
│   ├── __init__.py
│   ├── config.py
│   ├── data/
│   │   └── loader.py
│   ├── domain/
│   │   ├── costos.py
│   │   ├── tasas.py
│   │   ├── pd.py
│   │   ├── amortizacion.py
│   │   └── clv.py
│   └── optimization/
│       └── buscador_tea.py
├── api/
│   ├── __init__.py
│   ├── main.py
│   ├── schemas.py
│   └── router.py
├── tests/
│   ├── __init__.py
│   ├── test_clv.py
│   └── test_api.py
└── run.py"""

p = doc.add_paragraph()
run = p.add_run(code_block)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph('Requisitos de entrega:')
deliverables = [
    'Repositorio Git con commits incrementales (no un solo commit final).',
    'README.md con instrucciones de instalación y uso.',
    'requirements.txt o pyproject.toml.',
    'Tests actualizados que pasen correctamente.',
    'La API debe poder iniciarse con un solo comando.',
]
for d in deliverables:
    doc.add_paragraph(d, style='List Bullet')

doc.add_page_break()

# ── 7. Anexo ───────────────────────────────────────────
doc.add_heading('7. Anexo: Resumen del modelo actual', level=1)
doc.add_paragraph(
    'El modelo CLV calcula la Tasa Efectiva Anual (TEA) que hace que el Customer'
    ' Lifetime Value unitario sea igual al ROA target. A continuación se describe'
    ' el flujo de cálculo:'
)

doc.add_heading('Entradas', level=2)
inputs = [
    'Principal (P), plazo en meses (n), ROA target (ROA*)',
    'Costos: originación (c_orig) y mantenimiento anual (c_mant)',
    'Curva de tasas de transferencia (fondeo por plazo)',
    'Curva de probabilidad de default marginal (PD mensual)',
]
for inp in inputs:
    doc.add_paragraph(inp, style='List Bullet')

doc.add_heading('Modelo', level=2)
steps = [
    'Para un candidato TEA, se calcula TEM = (1+TEA)^(1/12) - 1',
    'Cuota francesa: C = P × TEM × (1+TEM)^n / ((1+TEM)^n - 1)',
    'Para cada mes t (1..n):',
    '   - Supervivencia: S(t) = Π(1 - PD_i) para i=1..t',
    '   - Saldo: saldo(t-1) según cronograma francés',
    '   - Costo fondeo: saldo × tasa_fondeo(t) / 12',
    '   - Costo mantenimiento: saldo × c_mant / 12',
    '   - Flujo neto esperado: (cuota - fondeo - mantenimiento) × S(t)',
    '   - Descuento: 1 / (1 + tasa_fondeo_mensual)^t',
    'PV = Σ flujos descontados',
    'CLV_unit = (-P - c_orig×P + PV) / P',
    'Optimización: encontrar TEA tal que CLV_unit ≈ ROA*',
]
for step in steps:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('Problemas conocidos del código actual', level=2)
issues = [
    'Algoritmo de búsqueda ineficiente: barrido lineal + bisección lenta',
    'iterrows() en todos los loops en vez de vectorización',
    'Supervivencia recalculada desde cero en cada mes (O(n²))',
    'Los CSVs se releen múltiples veces dentro de loops',
    '9 módulos con acoplamiento alto y namespaces contaminados',
    'Variables globales mutadas desde múltiples puntos',
    'Código duplicado (métodos copiados, cálculos repetidos 2-3 veces)',
]
for issue in issues:
    doc.add_paragraph(issue, style='List Bullet')

# ── Guardar ──────────────────────────────────────────────
ruta_salida = '/Users/masa/Library/CloudStorage/OneDrive-Credicorp/1. projects/recruting /test1/Examen_Tecnico_CLV.docx'
doc.save(ruta_salida)
print(f"Documento generado: {ruta_salida}")
